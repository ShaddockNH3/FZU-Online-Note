from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_OUT = DOCS / "generated" / "项目说明.docx"
PPTX_OUT = ROOT / "presentation" / "generated" / "答辩PPT.pptx"


def find_project_doc() -> Path:
    for path in DOCS.glob("*.md"):
        text = path.read_text(encoding="utf-8")
        if "FZU Online Note 项目说明" in text:
            return path
    raise FileNotFoundError("未找到项目说明 Markdown")


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F1EB")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(45, 45, 45)


def generate_docx():
    source = find_project_doc()
    text = source.read_text(encoding="utf-8")
    pages = [part.strip() for part in text.split("<!-- PAGE BREAK -->")]
    if len(pages) != 7:
        raise ValueError(f"项目说明应为 7 页，当前为 {len(pages)} 页")

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for page_index, page in enumerate(pages):
        if page_index:
            doc.add_page_break()

        in_code = False
        code_lines = []
        for raw_line in page.splitlines():
            line = raw_line.rstrip()
            if line.startswith("```"):
                if not in_code:
                    in_code = True
                    code_lines = []
                else:
                    add_code_block(doc, code_lines)
                    in_code = False
                continue
            if in_code:
                code_lines.append(line)
                continue
            if not line.strip():
                continue
            if line.startswith("# "):
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run(line[2:].strip())
                set_run_font(run, size=18, bold=True, color=(90, 61, 51))
            elif line.startswith("## "):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run(line[3:].strip())
                set_run_font(run, size=14, bold=True, color=(185, 68, 96))
            elif line.startswith("- "):
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                run = paragraph.add_run(line[2:].strip().replace("`", ""))
                set_run_font(run, size=10.2)
            else:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Pt(21)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.08
                run = paragraph.add_run(line.replace("`", ""))
                set_run_font(run, size=10.2)

    doc.save(DOCX_OUT)
    return DOCX_OUT


def xml_text(text):
    return escape(text).replace("\n", "&#10;")


def content_types():
    overrides = [
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
        '<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>',
        '<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/>',
        '<Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>',
        '<Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>',
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
    ]
    for index in range(1, 7):
        overrides.append(
            f'<Override PartName="/ppt/slides/slide{index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>'
        )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        + "".join(overrides)
        + "</Types>"
    )


def root_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>"""


def presentation_xml():
    slide_ids = "".join(
        f'<p:sldId id="{255 + index}" r:id="rId{index}"/>' for index in range(1, 7)
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId7"/></p:sldMasterIdLst>
  <p:sldIdLst>{slide_ids}</p:sldIdLst>
  <p:sldSz cx="12192000" cy="6858000" type="wide"/>
  <p:notesSz cx="6858000" cy="9144000"/>
</p:presentation>"""


def presentation_rels():
    rels = []
    for index in range(1, 7):
        rels.append(
            f'<Relationship Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" '
            f'Target="slides/slide{index}.xml"/>'
        )
    rels.append(
        '<Relationship Id="rId7" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" '
        'Target="slideMasters/slideMaster1.xml"/>'
    )
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(rels)
        + "</Relationships>"
    )


def slide_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>
</Relationships>"""


def shape_xml(shape_id, name, x, y, cx, cy, text, font_size, bold=False, color="1D2433"):
    paragraphs = []
    for line in text.split("\n"):
        paragraphs.append(
            f'<a:p><a:r><a:rPr lang="zh-CN" sz="{font_size}" b="{1 if bold else 0}">'
            f'<a:solidFill><a:srgbClr val="{color}"/></a:solidFill>'
            '</a:rPr><a:t>'
            f"{xml_text(line)}"
            "</a:t></a:r></a:p>"
        )
    return f"""
<p:sp>
  <p:nvSpPr><p:cNvPr id="{shape_id}" name="{name}"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr>
  <p:spPr><a:xfrm><a:off x="{x}" y="{y}"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>
  <p:txBody><a:bodyPr wrap="square" anchor="t"/><a:lstStyle/>{"".join(paragraphs)}</p:txBody>
</p:sp>"""


def slide_xml(title, bullets, footer):
    bullet_text = "\n".join(f"• {item}" for item in bullets)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:bg><p:bgPr><a:solidFill><a:srgbClr val="FFF7F4"/></a:solidFill><a:effectLst/></p:bgPr></p:bg>
    <p:spTree>
      <p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
      <p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>
      <p:sp>
        <p:nvSpPr><p:cNvPr id="2" name="Accent"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
        <p:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="12192000" cy="320000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:solidFill><a:srgbClr val="E96F8C"/></a:solidFill></p:spPr>
        <p:txBody><a:bodyPr/><a:lstStyle/><a:p/></p:txBody>
      </p:sp>
      {shape_xml(3, "Title", 740000, 680000, 10500000, 720000, title, 3400, True, "5A3D33")}
      {shape_xml(4, "Bullets", 960000, 1710000, 10100000, 3900000, bullet_text, 2100, False, "1D2433")}
      {shape_xml(5, "Footer", 760000, 6250000, 10600000, 280000, footer, 1100, False, "667085")}
    </p:spTree>
  </p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>"""


def slide_master_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld>
  <p:clrMap bg1="lt1" tx1="dk1" bg2="lt2" tx2="dk2" accent1="accent1" accent2="accent2" accent3="accent3" accent4="accent4" accent5="accent5" accent6="accent6" hlink="hlink" folHlink="folHlink"/>
  <p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst>
  <p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles>
</p:sldMaster>"""


def slide_master_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/>
</Relationships>"""


def slide_layout_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
 xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
 xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank" preserve="1">
  <p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr></p:spTree></p:cSld>
  <p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sldLayout>"""


def slide_layout_rels():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/>
</Relationships>"""


def theme_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="FZU Online Note">
  <a:themeElements>
    <a:clrScheme name="Sweet"><a:dk1><a:srgbClr val="1D2433"/></a:dk1><a:lt1><a:srgbClr val="FFFDF9"/></a:lt1><a:dk2><a:srgbClr val="5A3D33"/></a:dk2><a:lt2><a:srgbClr val="FFF7F4"/></a:lt2><a:accent1><a:srgbClr val="E96F8C"/></a:accent1><a:accent2><a:srgbClr val="7BCBB7"/></a:accent2><a:accent3><a:srgbClr val="F0BD59"/></a:accent3><a:accent4><a:srgbClr val="82A8DC"/></a:accent4><a:accent5><a:srgbClr val="8B7BC9"/></a:accent5><a:accent6><a:srgbClr val="667085"/></a:accent6><a:hlink><a:srgbClr val="B34460"/></a:hlink><a:folHlink><a:srgbClr val="8B7BC9"/></a:folHlink></a:clrScheme>
    <a:fontScheme name="Microsoft YaHei"><a:majorFont><a:latin typeface="Microsoft YaHei"/><a:ea typeface="微软雅黑"/></a:majorFont><a:minorFont><a:latin typeface="Microsoft YaHei"/><a:ea typeface="微软雅黑"/></a:minorFont></a:fontScheme>
    <a:fmtScheme name="Default"><a:fillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:fillStyleLst><a:lnStyleLst><a:ln w="6350"><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:ln></a:lnStyleLst><a:effectStyleLst><a:effectStyle><a:effectLst/></a:effectStyle></a:effectStyleLst><a:bgFillStyleLst><a:solidFill><a:schemeClr val="phClr"/></a:solidFill></a:bgFillStyleLst></a:fmtScheme>
  </a:themeElements>
</a:theme>"""


def app_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
 xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
  <Application>Microsoft PowerPoint</Application><Slides>6</Slides>
</Properties>"""


def core_xml():
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:dcterms="http://purl.org/dc/terms/"
 xmlns:dcmitype="http://purl.org/dc/dcmitype/"
 xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <dc:title>FZU Online Note 答辩PPT</dc:title>
  <dc:creator>待填写</dc:creator>
  <cp:lastModifiedBy>Codex</cp:lastModifiedBy>
</cp:coreProperties>"""


def generate_pptx():
    PPTX_OUT.parent.mkdir(parents=True, exist_ok=True)
    slides = [
        (
            "FZU Online Note",
            ["选题方向：在线笔记", "基础技术路线：HTML + CSS + JavaScript", "数据持久化：localStorage", "目标：完成可演示、可使用的轻量级笔记管理工具"],
        ),
        (
            "功能概览",
            ["新增、编辑、删除、复制笔记", "置顶、搜索、筛选、排序", "角色主题图库与数据概览", "刷新页面后笔记仍然保留"],
        ),
        (
            "课程要求对应",
            ["语义化 HTML：header / main / section / form / template", "CSS：Grid、Flex、响应式布局与视觉反馈", "JavaScript：表单校验、动态渲染、事件处理", "数据管理：完整 CRUD + localStorage"],
        ),
        (
            "源码结构",
            ["index.html：页面结构、表单、卡片模板", "styles.css：布局、主题色、响应式与状态样式", "app.js：数据模型、校验、渲染、CRUD、持久化", "README 与 docs：运行说明、需求文档、项目说明"],
        ),
        (
            "遇到的问题",
            ["docx 需求文档样式不完整：按编号和列表重新整理为 Markdown", "新增与编辑容易混淆：用隐藏 note-id 区分新增和更新", "搜索、筛选、排序逻辑分散：集中到 getVisibleNotes 处理", "角色图片外链可能不稳定：准备本地回退缓存并写入 .gitignore"],
        ),
        (
            "总结与改进",
            ["已完成基本要求：结构、样式、校验、动态渲染、持久化、CRUD", "项目运行简单：直接打开 index.html 即可", "后续可扩展 Markdown 预览、导入导出、标签管理", "进阶方向：Vue 重构或加入后端登录与云端同步"],
        ),
    ]
    with ZipFile(PPTX_OUT, "w", ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", content_types())
        package.writestr("_rels/.rels", root_rels())
        package.writestr("ppt/presentation.xml", presentation_xml())
        package.writestr("ppt/_rels/presentation.xml.rels", presentation_rels())
        package.writestr("ppt/slideMasters/slideMaster1.xml", slide_master_xml())
        package.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", slide_master_rels())
        package.writestr("ppt/slideLayouts/slideLayout1.xml", slide_layout_xml())
        package.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", slide_layout_rels())
        package.writestr("ppt/theme/theme1.xml", theme_xml())
        package.writestr("docProps/app.xml", app_xml())
        package.writestr("docProps/core.xml", core_xml())
        for index, (title, bullets) in enumerate(slides, start=1):
            package.writestr(
                f"ppt/slides/slide{index}.xml",
                slide_xml(title, bullets, f"FZU Online Note 答辩 - {index}/6"),
            )
            package.writestr(f"ppt/slides/_rels/slide{index}.xml.rels", slide_rels())
    return PPTX_OUT


if __name__ == "__main__":
    docx_path = generate_docx()
    print(f"generated: {docx_path}")
    print("PPT: run scripts/generate_presentation.ps1 to generate the PowerPoint file with Office COM.")
